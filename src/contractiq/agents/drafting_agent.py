from __future__ import annotations

import datetime
import logging
import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openai import OpenAI

from contractiq.agents.checklists import get_checklist
from contractiq.agents.models import (
    REVIEW_BANNER,
    CompletenessReport,
    DraftClause,
    DraftResult,
)
from contractiq.agents.vendor_resolution import extract_vendor_hint, resolve_vendor_doc_ids
from contractiq.config import settings
from contractiq.extraction.db import ContractRecord, get_session
from contractiq.extraction.models import ClauseType
from contractiq.retrieval.models import RetrievedChunk
from contractiq.retrieval.vector_store import dense_search

logger = logging.getLogger(__name__)

DRAFTS_DIR = Path("data/processed/drafts")
RAW_DIR = Path("data/raw")

NUMERIC_RE = re.compile(r"\$?\d[\d,]*(?:\.\d+)?%?")

SMOOTHING_SYSTEM_PROMPT = (
    "You lightly edit a single contract clause for a reusable draft template. You may ONLY:\n"
    '- Fix grammatical seams left by placeholder substitution (e.g. "The [VENDOR NAME] Inc. '
    'shall" -> "[VENDOR NAME] shall")\n'
    "- Normalize capitalization and spacing\n"
    "- Improve sentence flow\n\n"
    "You must NOT change, add, or remove any number, date, percentage, dollar amount, day "
    "count, or any substantive obligation, right, or condition. Return ONLY the edited clause "
    "text, with no commentary."
)


def _get_vendor_name(doc_id: str) -> str | None:
    session = get_session()
    record = session.query(ContractRecord).filter_by(doc_id=doc_id).one_or_none()
    session.close()
    return record.vendor if record else None


def _substitute_placeholders(text: str, vendor_name: str | None) -> str:
    if not vendor_name:
        return text
    return text.replace(vendor_name, "[VENDOR NAME]")


# Comparing raw numeric-token strings was rejecting smoothing on purely
# cosmetic differences ("07" -> "7", "1,000" spacing) that are common
# artifacts of OCR'd source text and carry no substantive change -- which
# meant the noisiest (most OCR-garbled) clauses were guaranteed to stay
# unsmoothed, since they're exactly the ones most likely to trip a naive
# string comparison. Stripping leading zeros/commas before comparing keeps
# the actual safety net (a genuine value change like "5 years" -> "3 years"
# still fails to match) while no longer flagging cosmetic reformatting.
def _normalize_numeric_token(token: str) -> str:
    has_dollar = token.startswith("$")
    has_percent = token.endswith("%")
    digits = token.strip("$%").replace(",", "")
    if "." in digits:
        int_part, frac_part = digits.split(".", 1)
        digits = f"{int_part.lstrip('0') or '0'}.{frac_part}"
    else:
        digits = digits.lstrip("0") or "0"
    return f"{'$' if has_dollar else ''}{digits}{'%' if has_percent else ''}"


def _extract_numeric_tokens(text: str) -> set[str]:
    return {_normalize_numeric_token(t) for t in NUMERIC_RE.findall(text)}


def _smooth_clause(text: str, client: OpenAI) -> tuple[str, bool]:
    """Returns (final_text, was_smoothed). Falls back to the unmodified
    input if the LLM's edit changes the set of numeric tokens present --
    this check, not the prompt instruction, is the actual safety net."""
    try:
        completion = client.chat.completions.create(
            model=settings.chat_model,
            messages=[
                {"role": "system", "content": SMOOTHING_SYSTEM_PROMPT},
                {"role": "user", "content": text},
            ],
            temperature=0,
        )
        smoothed = completion.choices[0].message.content or ""
    except Exception:
        logger.exception("Clause smoothing failed; using unmodified precedent text.")
        return text, False

    if _extract_numeric_tokens(smoothed) != _extract_numeric_tokens(text):
        logger.warning("Smoothing altered numeric content; reverting to raw precedent text.")
        return text, False

    return smoothed, True


def _build_where(clause_type: ClauseType, doc_ids: set[str] | None) -> dict:
    clauses: list[dict] = [{"clause_type": clause_type.value}]
    if doc_ids:
        clauses.append({"doc_id": {"$in": sorted(doc_ids)}})
    return clauses[0] if len(clauses) == 1 else {"$and": clauses}


def _find_precedent(
    clause_type: ClauseType,
    business_brief: str,
    agreement_type: str,
    client: OpenAI,
    doc_ids: set[str] | None = None,
) -> tuple[RetrievedChunk | None, bool | None]:
    """Returns (chunk, from_named_vendor). from_named_vendor is None when no
    vendor was named (doc_ids is None) -- vendor-scoping never applied.
    True/False otherwise: whether the match came from the named vendor's own
    document, or the search had to fall back to the best corpus-wide match
    because the named vendor has no precedent of this clause type."""
    query = f"{clause_type.value.replace('_', ' ')} clause for a {agreement_type} agreement. {business_brief}"

    if doc_ids:
        scoped = dense_search(query, client, top_k=1, where=_build_where(clause_type, doc_ids))
        if scoped:
            return scoped[0], True

    results = dense_search(query, client, top_k=1, where=_build_where(clause_type, None))
    if not results:
        return None, None
    return results[0], (False if doc_ids else None)


def _add_toc_field(doc: DocxDocument) -> None:
    """Inserts a real Word TOC field (not a static list) keyed off the
    document's Heading styles -- the same mechanism Word's own
    References > Table of Contents button produces. python-docx has no
    high-level API for this since Word, not python-docx, computes the
    pagination; the field shows placeholder text until the reader opens it
    in Word and updates it (right-click -> Update Field, or F9), which is
    the standard, expected way a programmatically generated Word TOC works."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose \"Update Field\" to generate the table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_begin)
    r.append(instr_text)
    r.append(fld_separate)
    r.append(placeholder)
    r.append(fld_end)


# Best-effort only: table structure is never preserved upstream of drafting
# (ingestion/extraction only ever recorded a boolean "does this page have a
# table" flag -- see extraction/recon.py -- not the table's actual rows), and
# ~85% of this corpus is scanned/OCR'd pages with no native table structure
# for PyMuPDF to find at all. So this re-opens the *original* source PDF at
# draft time and looks for a table on the precedent's exact page -- it will
# only ever succeed for the native (non-scanned) fraction of the corpus, and
# returns None (silently, no table in the draft) everywhere else.
def _extract_table_from_source(source_file: str, page_number: int | None) -> list[list[str]] | None:
    if not page_number:
        return None
    path = RAW_DIR / source_file
    if not path.exists() or path.suffix.lower() != ".pdf":
        return None
    try:
        with fitz.open(path) as pdf:
            if not (1 <= page_number <= len(pdf)):
                return None
            tables = pdf[page_number - 1].find_tables()
            if not tables.tables:
                return None
            data = tables.tables[0].extract()
    except Exception:
        logger.exception("Table extraction failed for %s p.%s", source_file, page_number)
        return None
    if not data or not data[0]:
        return None
    return data


def _add_table(doc: DocxDocument, data: list[list[str]]) -> None:
    n_cols = max(len(row) for row in data)
    table = doc.add_table(rows=len(data), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(data):
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) and row[j] is not None else ""
            table.cell(i, j).text = str(cell_text)


def _build_docx(
    agreement_type: str,
    business_brief: str,
    clauses: list[DraftClause],
    completeness: CompletenessReport,
    docx_path: Path,
) -> None:
    doc = DocxDocument()

    banner_para = doc.add_paragraph()
    banner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner_para.add_run(f"⚠ {REVIEW_BANNER} ⚠")
    banner_run.bold = True
    banner_run.font.size = Pt(14)
    banner_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_heading(f"{agreement_type} — Draft", level=1)
    doc.add_paragraph(f"Generated: {datetime.datetime.utcnow().isoformat()}Z")
    doc.add_paragraph("Business context (as provided, verbatim): " + business_brief)

    doc.add_heading("Table of Contents", level=2)
    _add_toc_field(doc)

    doc.add_heading("Completeness Report", level=2)
    if not completeness.agreement_type_recognized:
        note = doc.add_paragraph()
        note.add_run(
            f'Agreement type "{agreement_type}" was not recognized -- a generic checklist '
            "was used instead."
        ).italic = True
    for clause_type in completeness.required:
        status = "present" if clause_type in completeness.present else "MISSING"
        doc.add_paragraph(f"  - {clause_type.value}: {status}")

    doc.add_heading("Clauses", level=2)
    for clause in clauses:
        doc.add_heading(clause.clause_type.value.replace("_", " ").title(), level=3)
        if clause.status == "missing":
            p = doc.add_paragraph()
            p.add_run("⚠ MISSING — no precedent found for this clause type").italic = True
            continue

        doc.add_paragraph(clause.text)

        table_data = _extract_table_from_source(clause.source_document, clause.source_page)
        if table_data:
            _add_table(doc, table_data)

        source_p = doc.add_paragraph()
        source_text = f"Source: {clause.source_document}, {clause.source_section}, p.{clause.source_page}"
        if not clause.smoothed:
            source_text += " (raw precedent text -- automated smoothing was rejected)"
        if clause.from_named_vendor is False:
            source_text += " (not found in the named vendor's own contract -- using corpus precedent instead)"
        source_run = source_p.add_run(source_text)
        source_run.italic = True
        source_run.font.size = Pt(9)

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_run = footer_para.add_run(REVIEW_BANNER)
    footer_run.bold = True
    footer_run.font.size = Pt(8)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def drafting_agent(agreement_type: str, business_brief: str) -> DraftResult:
    """Assembles a draft from retrieved precedent clauses -- never
    generates clause language freely. The OpenAI API is used only to
    lightly smooth each selected precedent's phrasing (validated, see
    _smooth_clause), never to select clauses or invent missing ones.

    If the business brief names a specific vendor/contract (e.g. "based on
    the Asif & Co contract"), precedent search is scoped to that vendor's own
    documents first, falling back to the best corpus-wide match only for
    clause types that vendor's contract doesn't have -- mirrors the same
    fix applied to the RAG agent (agents/rag_agent.py) for the same
    boilerplate-corpus ambiguity problem."""
    if not settings.openai_api_key:
        raise RuntimeError(
            "OPENAI_API_KEY is not set. Copy .env.example to .env and fill it in "
            "before running the drafting agent."
        )
    client = OpenAI(api_key=settings.openai_api_key)

    checklist, recognized = get_checklist(agreement_type)

    vendor_hint = extract_vendor_hint(business_brief, client)
    doc_ids = resolve_vendor_doc_ids(vendor_hint) if vendor_hint else None
    if vendor_hint and not doc_ids:
        logger.warning("Drafting brief named a vendor (%r) but it didn't resolve to any document.", vendor_hint)

    clauses: list[DraftClause] = []
    for clause_type in checklist:
        precedent, from_named_vendor = _find_precedent(
            clause_type, business_brief, agreement_type, client, doc_ids=doc_ids
        )
        if precedent is None:
            clauses.append(DraftClause(clause_type=clause_type, status="missing"))
            continue

        vendor_name = _get_vendor_name(precedent.doc_id)
        prepared_text = _substitute_placeholders(precedent.text, vendor_name)
        final_text, smoothed = _smooth_clause(prepared_text, client)

        clauses.append(
            DraftClause(
                clause_type=clause_type,
                status="drafted",
                text=final_text,
                source_document=precedent.source_file,
                source_section=precedent.section_title
                or (f"Clause {precedent.clause_number}" if precedent.clause_number else None),
                source_page=precedent.page,
                smoothed=smoothed,
                from_named_vendor=from_named_vendor,
            )
        )

    present = [c.clause_type for c in clauses if c.status == "drafted"]
    missing = [c.clause_type for c in clauses if c.status == "missing"]
    completeness = CompletenessReport(
        agreement_type=agreement_type,
        agreement_type_recognized=recognized,
        required=checklist,
        present=present,
        missing=missing,
    )

    timestamp = datetime.datetime.utcnow().strftime("%Y%m%d%H%M%S")
    safe_type = re.sub(r"[^A-Za-z0-9_-]", "_", agreement_type)
    docx_path = DRAFTS_DIR / f"{safe_type}_{timestamp}.docx"
    _build_docx(agreement_type, business_brief, clauses, completeness, docx_path)

    return DraftResult(
        agreement_type=agreement_type,
        business_brief=business_brief,
        clauses=clauses,
        completeness=completeness,
        docx_path=str(docx_path),
    )
