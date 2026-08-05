from __future__ import annotations

import csv
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from contractiq.extraction.models import ReconResult
from contractiq.ingestion.loaders import SUPPORTED_SUFFIXES, load_document, needs_ocr

RAW_DIR = Path("data/raw")
PROCESSED_DIR = Path("data/processed")


def _scan_pdf(path: Path) -> ReconResult:
    pages_native = 0
    pages_scanned = 0
    char_count = 0
    likely_tables = False

    with fitz.open(path) as pdf:
        page_count = pdf.page_count
        for page in pdf:
            text = page.get_text().strip()
            char_count += len(text)
            if needs_ocr(text):
                pages_scanned += 1
            else:
                pages_native += 1
            if not likely_tables and page.find_tables().tables:
                likely_tables = True

    if pages_scanned == 0:
        classification = "native"
    elif pages_native == 0:
        classification = "scanned"
    else:
        classification = "mixed"

    return ReconResult(
        file=path.name,
        file_format="pdf",
        page_count=page_count,
        char_count=char_count,
        classification=classification,
        pages_native=pages_native,
        pages_scanned=pages_scanned,
        likely_tables=likely_tables,
    )


def _scan_docx(path: Path) -> ReconResult:
    document = load_document(path)
    char_count = sum(len(p.text) for p in document.pages)
    likely_tables = len(DocxDocument(path).tables) > 0

    return ReconResult(
        file=path.name,
        file_format="docx",
        page_count=None,
        char_count=char_count,
        classification="native",
        pages_native=1,
        pages_scanned=0,
        likely_tables=likely_tables,
    )


def scan_directory(input_dir: Path = RAW_DIR) -> list[ReconResult]:
    results: list[ReconResult] = []
    for path in sorted(input_dir.iterdir()):
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            continue
        results.append(_scan_pdf(path) if suffix == ".pdf" else _scan_docx(path))
    return results


def write_csv(results: list[ReconResult], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(ReconResult.model_fields))
        writer.writeheader()
        for r in results:
            writer.writerow(r.model_dump())


def print_summary(results: list[ReconResult]) -> None:
    native = sum(1 for r in results if r.classification == "native")
    scanned = sum(1 for r in results if r.classification == "scanned")
    mixed = sum(1 for r in results if r.classification == "mixed")
    tables = sum(1 for r in results if r.likely_tables)

    header = f"{'File':<40} {'Format':<6} {'Pages':<7} {'Chars':<8} {'Class':<9} {'Tables':<6}"
    print(header)
    print("-" * len(header))
    for r in results:
        pages = r.page_count if r.page_count is not None else "N/A"
        print(
            f"{r.file:<40} {r.file_format:<6} {str(pages):<7} {r.char_count:<8} "
            f"{r.classification:<9} {'yes' if r.likely_tables else 'no':<6}"
        )

    print()
    print(f"Total files: {len(results)}")
    print(f"  native:  {native}")
    print(f"  scanned: {scanned}")
    print(f"  mixed:   {mixed}")
    print(f"  likely-tables: {tables}")


def run(
    input_dir: Path = RAW_DIR,
    output_csv: Path = PROCESSED_DIR / "recon_report.csv",
) -> list[ReconResult]:
    results = scan_directory(input_dir)
    write_csv(results, output_csv)
    print_summary(results)
    return results


if __name__ == "__main__":
    run()
