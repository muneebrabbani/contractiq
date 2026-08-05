from __future__ import annotations

import hashlib
from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from docx import Document as DocxDocument
from PIL import Image

from contractiq.ingestion.models import Page, SourceDocument

# Pages with less extractable text than this are treated as scanned images
# and routed through OCR instead.
OCR_MIN_CHARS = 20

SUPPORTED_SUFFIXES = {".pdf", ".docx"}


def needs_ocr(text: str) -> bool:
    return len(text.strip()) < OCR_MIN_CHARS


def read_pdf_native_pages(path: Path) -> list[tuple[int, str]]:
    """Native (non-OCR) per-page text. Used where OCR must stay out of scope."""
    with fitz.open(path) as pdf:
        return [(i, page.get_text().strip()) for i, page in enumerate(pdf, start=1)]


def _load_pdf(path: Path) -> list[Page]:
    pages: list[Page] = []
    with fitz.open(path) as pdf:
        for i, page in enumerate(pdf, start=1):
            text = page.get_text().strip()
            ocr = False
            if needs_ocr(text):
                pix = page.get_pixmap(dpi=300)
                image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                text = pytesseract.image_to_string(image).strip()
                ocr = True
            pages.append(Page(page_number=i, text=text, ocr=ocr))
    return pages


def _load_docx(path: Path) -> list[Page]:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return [Page(page_number=1, text="\n".join(parts), ocr=False)]


def compute_doc_id(path: Path) -> str:
    digest = hashlib.sha256(path.read_bytes()).hexdigest()[:10]
    return f"{path.stem}-{digest}"


def load_document(path: Path) -> SourceDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages = _load_pdf(path)
        file_format = "pdf"
    elif suffix == ".docx":
        pages = _load_docx(path)
        file_format = "docx"
    else:
        raise ValueError(f"Unsupported file format: {suffix}")

    return SourceDocument(
        doc_id=compute_doc_id(path),
        file_path=str(path),
        file_format=file_format,
        pages=pages,
    )
